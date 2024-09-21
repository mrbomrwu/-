from matplotlib import pyplot as plt
import matplotlib.patches as patches
from matplotlib.offsetbox import TextArea, VPacker, AnnotationBbox
from docx import Document
from docx.shared import Inches
import io

# Function to draw the Genetic Algorithm Flowchart
def draw_ga_flowchart():
    fig, ax = plt.subplots(figsize=(8, 10))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 14)
    ax.axis('off')

    # Create process boxes
    def add_box(x, y, text, width=4, height=1):
        box = patches.FancyBboxPatch((x, y), width, height, boxstyle="round,pad=0.1", edgecolor='black', facecolor='lightblue')
        ax.add_patch(box)
        txt = ax.text(x + width / 2, y + height / 2, text, va='center', ha='center', fontsize=10, wrap=True)
        return box, txt

    # Create arrows
    def add_arrow(x, y, dx, dy):
        arrow = patches.FancyArrowPatch((x, y), (x + dx, y + dy), mutation_scale=10, color='black')
        ax.add_patch(arrow)

    # Add boxes and arrows for Genetic Algorithm
    y_start = 12
    steps = [
        "Start",
        "Initialize Population",
        "Evaluate Fitness of Population",
        "Selection",
        "Crossover",
        "Mutation",
        "Evaluate Fitness of New Population",
        "Termination Criteria Met?",
        "End"
    ]

    for i, step in enumerate(steps):
        add_box(3, y_start - i * 1.5, step)
        if i < len(steps) - 1:
            add_arrow(5, y_start - i * 1.5 - 0.1, 0, -1.3)

    # Add decision box for termination criteria
    box = patches.FancyBboxPatch((3, y_start - 10.5), 4, 1.5, boxstyle="round,pad=0.1", edgecolor='black', facecolor='lightyellow')
    ax.add_patch(box)
    txt = ax.text(5, y_start - 9.75, "Termination Criteria Met?", va='center', ha='center', fontsize=10, wrap=True)

    # Add Yes/No paths
    add_arrow(5, y_start - 11.5, -4, -1)
    ax.text(0.7, y_start - 12.1, "Yes", fontsize=10, va='center', ha='center')

    add_arrow(5, y_start - 11.5, 0, -1.3)
    ax.text(7.5, y_start - 12.1, "No", fontsize=10, va='center', ha='center')

    return fig

# Draw the flowchart and save it as an image
fig = draw_ga_flowchart()
image_stream = io.BytesIO()
fig.savefig(image_stream, format='png', bbox_inches='tight')
plt.close(fig)

# Create a Word document and insert the flowchart image
doc = Document()
doc.add_heading('Genetic Algorithm Flowchart', level=1)
doc.add_paragraph("The flowchart below illustrates the steps involved in a Genetic Algorithm:")

# Insert the image
doc.add_picture(image_stream, width=Inches(5.5))

# Save the document
doc_path = "genetic_algorithm_flowchart.docx"

doc.save(doc_path)

doc_path
